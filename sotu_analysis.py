#!/usr/bin/env python3
"""
sotu_analysis.py

Reads a speech text file and outputs:
 - Word count
 - Character count
 - Average word length
 - Average sentence length (words per sentence)
 - Word frequency distribution (sorted)
 - Top ten longest words

Also (optionally) creates a Word document report containing pseudocode, source,
and a snapshot of the textual output. Requires 'python-docx' to build the Word file.
"""

from collections import Counter
import re
import argparse
import textwrap
import os
import sys

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# --------------------------
# Utilities
# --------------------------
SENTENCE_END_RE = re.compile(r'(?<=[.!?])\s+')
WORD_TOKEN_RE = re.compile(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)*")  # keeps internal apostrophes/hyphens

def read_file(path, encoding='utf-8'):
    with open(path, 'r', encoding=encoding, errors='ignore') as f:
        return f.read()

def split_sentences(text):
    # Basic sentence splitting: split at . ! ? followed by whitespace
    # Trim whitespace and ignore empty segments
    parts = SENTENCE_END_RE.split(text.strip())
    parts = [p.strip() for p in parts if p.strip()]
    return parts

def tokenize_words(text):
    # Finds word tokens using WORD_TOKEN_RE
    return WORD_TOKEN_RE.findall(text)

def normalize_word(w):
    return w.lower()

def compute_statistics(text):
    # Character count: include all characters (including whitespace)
    char_count = len(text)

    # Sentence split
    sentences = split_sentences(text)
    sentence_count = len(sentences) if sentences else 0

    # Word tokens
    raw_words = tokenize_words(text)
    word_count = len(raw_words)

    # Normalized words for frequency
    norm_words = [normalize_word(w) for w in raw_words]

    # Average word length (characters in words / number of words)
    avg_word_length = (sum(len(w) for w in raw_words) / word_count) if word_count else 0.0

    # Average sentence length in words
    if sentence_count:
        words_per_sentence = [len(tokenize_words(s)) for s in sentences]
        avg_sentence_length = sum(words_per_sentence) / sentence_count
    else:
        avg_sentence_length = float(word_count)

    # Word frequency distribution (descending frequency, then alphabetic)
    freq = Counter(norm_words)
    freq_sorted = sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))

    # Top ten longest unique words (by length, then alphabetic)
    unique_words = set(raw_words)
    top_longest = sorted(unique_words, key=lambda w: (-len(w), w.lower()))[:10]

    return {
        'char_count': char_count,
        'word_count': word_count,
        'avg_word_length': avg_word_length,
        'avg_sentence_length': avg_sentence_length,
        'freq_sorted': freq_sorted,
        'top_longest': top_longest,
        'sentences': sentences,
        'raw_words': raw_words
    }

def print_table(stats, top_n_freq=20):
    print("\n=== Speech Analysis ===\n")
    # Summary table
    print(f"{'Metric':<30}{'Value'}")
    print("-" * 50)
    print(f"{'Word count':<30}{stats['word_count']}")
    print(f"{'Character count (incl. whitespace)':<30}{stats['char_count']}")
    print(f"{'Average word length (chars)':<30}{stats['avg_word_length']:.2f}")
    print(f"{'Average sentence length (words)':<30}{stats['avg_sentence_length']:.2f}")
    print()

    # Word frequency table (top N)
    print(f"Top {top_n_freq} words by frequency:")
    print(f"{'Rank':<5}{'Word':<20}{'Count'}")
    print("-" * 40)
    for i, (word, cnt) in enumerate(stats['freq_sorted'][:top_n_freq], start=1):
        print(f"{i:<5}{word:<20}{cnt}")
    print()

    # Top ten longest words
    print("Top 10 longest unique words:")
    for i, w in enumerate(stats['top_longest'], start=1):
        print(f"{i}. {w} ({len(w)} chars)")
    print()

def make_docx_report(input_path, stats, output_path='sotu_report.docx'):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

    doc = Document()
    doc.add_heading('State of the Union Speech Analysis', level=1)

    doc.add_heading('Input file', level=2)
    doc.add_paragraph(input_path)

    doc.add_heading('Summary', level=2)
    doc.add_paragraph(f"Word count: {stats['word_count']}")
    doc.add_paragraph(f"Character count (incl. whitespace): {stats['char_count']}")
    doc.add_paragraph(f"Average word length (chars): {stats['avg_word_length']:.2f}")
    doc.add_paragraph(f"Average sentence length (words): {stats['avg_sentence_length']:.2f}")

    doc.add_heading('Top words (frequency)', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Word'
    hdr_cells[2].text = 'Count'
    for i, (word, cnt) in enumerate(stats['freq_sorted'][:50], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = word
        row_cells[2].text = str(cnt)

    doc.add_heading('Top 10 Longest Words', level=2)
    for i, w in enumerate(stats['top_longest'], start=1):
        doc.add_paragraph(f"{i}. {w} ({len(w)} chars)")

    doc.add_heading('Pseudocode', level=2)
    pseudocode = textwrap.dedent("""\
    BEGIN
      INPUT path to speech file
      READ file content
      SPLIT text into sentences
      TOKENIZE words, normalize
      COMPUTE counts, averages, frequencies, longest words
      OUTPUT results
    END
    """)
    doc.add_paragraph(pseudocode)

    doc.add_heading('Source code', level=2)
    # read this source file and add as preformatted
    try:
        with open(__file__, 'r', encoding='utf-8') as f:
            source = f.read()
    except Exception:
        source = "Source code could not be embedded automatically."

    # add code in a monospaced-like paragraph (docx has limited style control)
    doc.add_paragraph(source)

    # Save
    doc.save(output_path)
    return output_path

# --------------------------
# CLI
# --------------------------
def main():
    parser = argparse.ArgumentParser(description="Analyze a State-of-the-Union style speech text file.")
    parser.add_argument('input', nargs='?', default='state_of_union.txt', help='Path to speech text file (default: state_of_union.txt)')
    parser.add_argument('--docx', action='store_true', help='Create a Word (.docx) report (requires python-docx)')
    parser.add_argument('--topn', type=int, default=20, help='How many top frequency words to display')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: input file '{args.input}' not found.", file=sys.stderr)
        sys.exit(2)

    text = read_file(args.input)
    stats = compute_statistics(text)

    print_table(stats, top_n_freq=args.topn)

    if args.docx:
        if not DOCX_AVAILABLE:
            print("python-docx is not installed. To create a Word report, run: pip install python-docx", file=sys.stderr)
            sys.exit(3)
        out = make_docx_report(args.input, stats, output_path='sotu_report.docx')
        print(f"Word report saved to: {out}")

if __name__ == '__main__':
    main()
